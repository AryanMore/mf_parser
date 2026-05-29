import os
from docx import Document
from models.schema import CobolProgram
from analyzer.semantic_engine import SemanticEngine


class DOCXGenerator:
    """
    Generates enterprise COBOL comprehension DOCX.
    Supports:
    - deterministic parser output
    - validated LLM sections
    """

    def __init__(
        self,
        output_dir: str = "output"
    ):
        self.semantic_engine = SemanticEngine()
        self.output_dir = output_dir
        os.makedirs(
            self.output_dir,
            exist_ok=True
        )

    # =====================================================
    # PUBLIC ENTRY
    # =====================================================
    def generate(
        self,
        program: CobolProgram
    ):
        program_name = (
            program.metadata.program_id
            if program.metadata.program_id
            else "unknown_program"
        )

        output_path = os.path.join(
            self.output_dir,
            f"{program_name}.docx"
        )

        doc = Document()

        # -----------------------------------
        # TITLE
        # -----------------------------------
        doc.add_heading(
            f"COBOL Program Comprehension - {program_name}",
            level=0
        )

        # =====================================================
        # 1. INTRODUCTION
        # =====================================================
        enabled = getattr(
            program,
            "enabled_sections",
            {}
        )

        llm = getattr(
            program,
            "llm_sections",
            {}
        )

        if enabled.get("introduction"):
            doc.add_heading(
                "1. Introduction",
                level=1
            )

            if enabled.get(
                "program_overview"
            ):
                doc.add_heading(
                    "1.1 Program Overview",
                    level=2
                )
                overview = (
                llm.get(
                    "program_overview"
                )
                or self.semantic_engine.generate_program_overview(program)
                )

                doc.add_paragraph(overview)

            if enabled.get(
                "objectives"
            ):
                doc.add_heading(
                    "1.2 Objectives",
                    level=2
                )
                objectives = (
                    llm.get("objectives")
                    or llm.get("objectives_refinement")
                    or self.semantic_engine.generate_objectives_summary(program)
                )

                doc.add_paragraph(objectives)

            if enabled.get(
                "scope"
            ):
                doc.add_heading(
                    "1.3 Scope",
                    level=2
                )
                scope = (
                    llm.get("scope")
                    or self.semantic_engine.generate_scope_summary(program)
                )

                doc.add_paragraph(scope)

            if enabled.get(
                "assumptions_constraints"
            ):
                doc.add_heading(
                    "1.4 Assumptions and Constraints",
                    level=2
                )
                constraints = (
                    llm.get(
                        "assumptions_constraints"
                    )
                    or self.semantic_engine.generate_constraints_summary(program)
                )

                doc.add_paragraph(constraints)
        # =====================================================
        # 2. DATABASE DETAILS
        # =====================================================
        if (
            enabled.get("database_details")
            or enabled.get("ims_segments")
            or enabled.get("idms_records")
        ):
            doc.add_heading("2. Database Details", level=1)
            doc.add_paragraph(
                self.semantic_engine.generate_database_summary(program)
            )

            if enabled.get(
                "db2_tables"
            ):
                doc.add_heading("2.1 DB2 Tables", level=2)

                if program.sql_blocks:
                    for sql in program.sql_blocks:
                        for table in sql.tables:
                            doc.add_paragraph(
                                table,
                                style="List Bullet"
                            )
                else:
                    doc.add_paragraph(
                        "No DB2 usage detected."
                    )
            else:
                doc.add_heading("2.1 DB2 Tables", level=2)
                doc.add_paragraph("No DB2 usage detected.")

            doc.add_heading("2.2 IMS Segments", level=2)
            doc.add_paragraph(
                "No IMS segments detected from parser-visible source."
            )

            doc.add_heading("2.3 IDMS Records", level=2)
            doc.add_paragraph(
                "No IDMS records detected from parser-visible source."
            )

        # =====================================================
        # 3. SYSTEM ARCHITECTURE
        # =====================================================
        if enabled.get(
            "system_architecture"
        ):
            doc.add_heading(
                "3. System Architecture",
                level=1
            )

            architecture = (
                llm.get(
                    "architecture_refinement"
                )
                or self.semantic_engine.generate_architecture_summary(program)
            )

            doc.add_paragraph(
                architecture
            )

            # 3.1 Components

            component_text = (
                self.semantic_engine
                .generate_component_diagram_text(program)
            )

            doc.add_paragraph(component_text)

            if enabled.get(
                "component_diagram"
            ):
                doc.add_heading("3.1 Component Diagram", level=2)

                components = set()

                components.update(
                    program.copybooks
                )
                components.update(
                    program.calls
                )

                for comp in sorted(
                    components
                ):
                    doc.add_paragraph(
                        comp,
                        style="List Bullet"
                    )

            # 3.2 Control flow
            if enabled.get(
                "control_flow_diagram"
            ):
                flow_summary = (
                    self.semantic_engine
                    .generate_control_flow_summary(program)
                )

                doc.add_paragraph(flow_summary)

                doc.add_heading("3.2 Control Flow Diagram", level=2)

                # Flow summary
                doc.add_paragraph(
                    f"PERFORM count: {program.control_flow.perform_count}"
                )
                doc.add_paragraph(
                    f"CALL count: {program.control_flow.call_count}"
                )
                doc.add_paragraph(
                    f"IF count: {program.control_flow.if_count}"
                )
                doc.add_paragraph(
                    f"EVALUATE count: {program.control_flow.evaluate_count}"
                )
                doc.add_paragraph(
                    f"Complexity score: {program.control_flow.complexity_score}"
                )

                # Flow edges
                for edge in program.control_flow.edges:
                    doc.add_paragraph(
                        f"{edge.source} → {edge.target} ({edge.relation})",
                        style="List Bullet"
                    )

        # =====================================================
        # 4. DETAILED DESIGN
        # =====================================================
        if enabled.get(
            "detailed_design"
        ):
            doc.add_heading(
                "4. Detailed Design",
                level=1
            )

            # -----------------------------------
            # 4.1 Program Structure
            # -----------------------------------
            if enabled.get(
                "program_structure"
            ):
                doc.add_heading(
                    "4.1 Program Structure",
                    level=2
                )

                structure = (
                    llm.get(
                        "structure_refinement"
                    )
                    or self.semantic_engine.generate_program_structure_summary(program)
                )

                doc.add_paragraph(
                    structure
                )

                for section in program.sections:
                    doc.add_paragraph(
                        section.name,
                        style="List Bullet"
                    )

            # -----------------------------------
            # 4.2 Algorithms
            # -----------------------------------
            if enabled.get(
                "algorithms"
            ):
                doc.add_heading(
                    "4.2 Algorithms",
                    level=2
                )
                algorithm_summary = (
                    llm.get(
                        "algorithm_refinement"
                    )
                    or self.semantic_engine.generate_algorithm_summary(program)
                )

                doc.add_paragraph(
                    algorithm_summary
                )

            # -----------------------------------
            # 4.3 Input/Output Specifications
            # -----------------------------------
            if enabled.get(
                "io_specifications"
            ):
                doc.add_heading(
                    "4.3 Input/Output Specifications",
                    level=2
                )

                io_summary = (
                    self.semantic_engine
                    .generate_io_summary(program)
                )

                doc.add_paragraph(io_summary)

                io_table = doc.add_table(
                    rows=1,
                    cols=2
                )
                io_table.style = (
                    "Table Grid"
                )

                io_table.cell(
                    0, 0
                ).text = "Operation"
                io_table.cell(
                    0, 1
                ).text = "Target"

                for op in (
                    program.file_operations
                ):
                    row = (
                        io_table
                        .add_row()
                        .cells
                    )

                    row[0].text = (
                        op.operation
                    )
                    row[1].text = (
                        op.target
                    )

        # =====================================================
        # 5. ADVANCED COMPREHENSION (LLM V2)
        # =====================================================
        llm = getattr(
            program,
            "llm_sections",
            {}
        )

        if llm:
            doc.add_heading(
                "5. Advanced Comprehension",
                level=1
            )

            # -----------------------------------
            # 5.1 Business Logic Inference
            # -----------------------------------
            business_logic = llm.get(
                "business_logic"
            )

            if business_logic:
                doc.add_heading(
                    "5.1 Business Logic Inference",
                    level=2
                )
                doc.add_paragraph(
                    business_logic
                )

            # -----------------------------------
            # 5.2 Paragraph Explanations
            # -----------------------------------
            paragraph_explanations = llm.get(
                "paragraph_explanations"
            )

            if paragraph_explanations:
                doc.add_heading(
                    "5.2 Paragraph Explanations",
                    level=2
                )

                for line in (
                    paragraph_explanations
                    .splitlines()
                ):
                    line = line.strip()

                    if not line:
                        continue

                    doc.add_paragraph(
                        line,
                        style="List Bullet"
                    )

            # -----------------------------------
            # 5.3 Program Walkthrough
            # -----------------------------------
            walkthrough = llm.get(
                "program_walkthrough"
            )

            if walkthrough:
                doc.add_heading(
                    "5.3 Program Walkthrough",
                    level=2
                )

                for line in (
                    walkthrough
                    .splitlines()
                ):
                    line = line.strip()

                    if not line:
                        continue

                    doc.add_paragraph(
                        line
                    )

            # -----------------------------------
            # 5.4 Humanized Technical Summary
            # -----------------------------------
            humanized = (
                llm.get(
                    "architecture_refinement"
                )
            )

            if humanized:
                doc.add_heading(
                    "5.4 Humanized Technical Summary",
                    level=2
                )
                doc.add_paragraph(
                    humanized
                )

        # =====================================================
        # APPENDIX / STRUCTURED FACTS
        # =====================================================
        doc.add_heading(
            "Appendix A - Parsed Metadata",
            level=1
        )

        metadata = doc.add_table(
            rows=3,
            cols=2
        )
        metadata.style = (
            "Table Grid"
        )

        metadata.cell(
            0, 0
        ).text = "Program ID"
        metadata.cell(
            0, 1
        ).text = (
            program.metadata.program_id
            or ""
        )

        metadata.cell(
            1, 0
        ).text = "Author"
        metadata.cell(
            1, 1
        ).text = (
            program.metadata.author
            or ""
        )

        metadata.cell(
            2, 0
        ).text = "Date Written"
        metadata.cell(
            2, 1
        ).text = (
            program.metadata
            .date_written
            or ""
        )

        doc.save(
            output_path
        )
        return output_path
